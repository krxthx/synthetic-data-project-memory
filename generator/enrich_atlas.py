from pathlib import Path
import json, csv
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors

ROOT=Path('data/atlas')
ART=ROOT/'artifacts'
ART.mkdir(exist_ok=True)

people=json.loads((ROOT/'people.json').read_text())
by_role={p['role']:p for p in people}
N=[p['name'] for p in people]

DECISIONS=[
('Phase 1 scope','Enterprise onboarding is phase 1; SMB is deferred','Samir Rao','Protect the launch window and avoid mixing materially different operating models.'),
('Identity authority','CRM remains authoritative for customer identity and onboarding status','Adrian Shah','Avoid duplicate master data and preserve existing customer-operations controls.'),
('Authentication','Atlas uses the Harbor interim SSO bridge for pilot','Rhea Pillai','Decouple the Atlas pilot from the full Harbor migration while retaining enterprise identity controls.'),
('Exception handling','High-risk onboarding exceptions require human approval','Sofia Green','Legal, security, and contractual exceptions require accountable human judgment.'),
('Regional process','A common stage taxonomy is mandatory; regions may add governed substeps','Rachel Brooks','Provide enterprise reporting consistency without erasing legitimate regional process differences.')]
RISKS=['duplicate customer records','regional stage mismatch','Harbor SSO timing','legal exception review','client readiness variance']

# Expand short document-like JSONL records into realistic source documents.
def expand(rec, source):
    m=rec['metadata']; typ=rec['type']; idx=int(rec['id'].split('-')[-1]); d=DECISIONS[idx%len(DECISIONS)]; risk=RISKS[idx%len(RISKS)]
    author=rec['author']; version=m.get('version','1.0')
    if source=='sharepoint':
        title={'business_announcement':'Atlas Program Update','status_report':'Atlas Weekly Delivery Status','steering_decision':'Atlas Steering Decision Record','client_update':'Atlas Client Readiness Update'}.get(typ,'Atlas Project Document')
        body=f'''{title}\nVersion: {version}\nOwner: {author}\nProject: Atlas Customer Onboarding\n\n1. Purpose and Context\nAtlas is modernizing the path from signed enterprise deal to production handoff. The program is intended to replace fragmented handoffs, inconsistent regional stage definitions, and opaque exception handling with a governed onboarding model. This document is a point-in-time project artifact and should be interpreted together with steering minutes, architecture decisions, and implementation specifications.\n\n2. Current Position\nThe current authoritative direction for {d[0].lower()} is: {d[1]}. This position was selected because {d[3]} The team has reviewed the impact across product, architecture, QA, client operations, and change management. Where older documents conflict with this position, they remain useful only as historical rationale.\n\n3. Business and Client Impact\nThe immediate business objective is to reduce onboarding ambiguity without creating a second system of record. Client stakeholders need predictable stage visibility, clear ownership of blocked tasks, and an escalation route when an onboarding case cannot follow the standard path. The client product owner has asked that regional flexibility remain possible, but not at the expense of enterprise reporting or auditability.\n\n4. Delivery Status\nThe primary risk under review is {risk}. The team is treating this as a delivery and operating-model concern rather than a purely technical defect. BA and product owners are validating requirements, architecture is checking downstream dependencies, and QA is maintaining acceptance coverage for standard and exception paths. No milestone change should be inferred unless a steering record explicitly approves it.\n\n5. Dependencies\nAtlas depends on Harbor for the interim SSO bridge. It also shares CRM identity-quality concerns with Nova. These dependencies are tracked separately because their timelines and owners differ from Atlas. A dependency delay does not automatically change Atlas scope; the project manager must first evaluate resequencing options.\n\n6. Decision and Rationale\nDecision: {d[1]}. Decision owner: {d[2]}. Rationale: {d[3]} Alternatives considered included retaining the legacy regional approach, waiting for all enterprise dependencies to complete, and allowing local teams to resolve the issue independently. Those alternatives were rejected because they either increase operational fragmentation or make the pilot timeline dependent on unrelated transformation work.\n\n7. Risks, Controls, and Open Questions\nThe current risk is {risk}. Controls include named ownership, explicit exception paths, audit logging for material decisions, and client-facing acceptance criteria. Open questions should be taken to the relevant SME rather than resolved through undocumented local convention.\n\n8. Actions\n- Business Analysis: update requirements and traceability where this decision changes acceptance criteria.\n- Architecture: confirm dependency and integration implications.\n- QA: maintain positive, negative, regression, and exception-path coverage.\n- Project Management: reflect material changes in steering status.\n- Client Product Owner: confirm that the operating outcome remains acceptable.\n\n9. Historical Note\nEarlier working assumptions are intentionally retained in the corpus. They are not current truth when superseded by a later accepted steering or architecture decision. The memory layer should preserve both the latest position and the reason the earlier approach existed.'''
    elif source=='onedrive':
        title={'implementation_spec':'Atlas Implementation Specification','working_notes':'Atlas Working Design Notes','test_plan':'Atlas QA and Acceptance Plan','process_map':'Atlas Onboarding Process Definition'}.get(typ,'Atlas Working Document')
        body=f'''{title}\nVersion: {version}\nAuthor: {author}\nStatus: Working project artifact\n\n1. Objective\nThis document describes implementation-level detail for Atlas Customer Onboarding. It is intentionally more detailed than the executive SharePoint material and may contain working assumptions that are later superseded. The current topic is {d[0].lower()}.\n\n2. Requirement\nThe implementation must support the following current direction: {d[1]}. The design must preserve CRM as the authoritative identity/status source, expose a clear onboarding stage to internal and client-facing users, and prevent local workflow customizations from silently changing enterprise semantics.\n\n3. Functional Flow\nA signed enterprise customer enters Atlas after the commercial handoff is complete. Atlas resolves the CRM customer record, initializes the governed onboarding stage model, assigns required tasks, and exposes status to authorized project participants. Standard tasks proceed automatically when prerequisites are met. High-risk exceptions pause the affected path and require an accountable human approval before the workflow continues.\n\n4. Integration Behavior\nAuthentication for the pilot uses the Harbor interim SSO bridge. CRM identifiers are carried through Atlas rather than recreated. Notifications are derived from onboarding state transitions. Integration failures must be retryable and observable; they must not create a second customer record or silently advance a stage.\n\n5. Data and State\nCore state includes customer identifier, onboarding case identifier, enterprise stage, optional regional substep, task ownership, exception state, approval evidence, timestamps, and source references. Historical transitions are retained for audit and support. Free-form notes may supplement but cannot override governed state.\n\n6. Current Constraint\nThe primary constraint for this revision is {risk}. The implementation should fail visibly where the constraint prevents safe progression. Temporary workarounds require an owner, expiry condition, and reference to the decision that authorized them.\n\n7. Validation\nQA must cover normal progression, duplicate identity detection, SSO failure/recovery, regional substeps, exception approval/rejection, stale client sessions, retries, and audit history. Client acceptance must verify both operational usability and status accuracy.\n\n8. Ownership\nDocument owner: {author}. Decision owner for this topic: {d[2]}. Solution architecture is owned by Sofia Green with enterprise dependency review from Rhea Pillai and Marcus Green. Rachel Chen owns release-quality interpretation.\n\n9. Decision Trace\nCurrent decision: {d[1]}. Rationale: {d[3]} Older approaches may still appear in meeting notes and prior versions. They should be treated as historical unless reaffirmed in a later decision record.\n\n10. Open Items\n- Confirm client acceptance wording for exception states.\n- Validate Harbor bridge operational support before pilot cutover.\n- Reconcile regional terminology with the enterprise stage dictionary.\n- Confirm duplicate-record remediation ownership with CRM operations.\n- Ensure training material matches the final process rather than early prototypes.'''
    else:
        title='Atlas Architecture / Implementation Decision'
        body=f'''{title}\nTopic: {d[0]}\nStatus: Accepted unless superseded\nOwner: {author}\n\nContext\nAtlas must provide a consistent enterprise onboarding model while integrating with existing CRM identity, Harbor authentication, regional operating processes, and client-facing status expectations. The design has to remain supportable after the pilot, not merely demonstrate a happy path. The current pressure point is {risk}.\n\nDecision\n{d[1]}.\n\nRationale\n{d[3]} The team explicitly values traceability and operational clarity over locally optimized shortcuts. The chosen direction also keeps ownership boundaries clear between Atlas, CRM operations, and Harbor.\n\nArchitecture Consequences\nAtlas owns onboarding orchestration and governed workflow state. CRM owns customer identity and authoritative commercial/onboarding status fields that are explicitly synchronized. Harbor owns authentication and identity controls. Regional teams may extend the process only through governed substeps. High-risk exception decisions are represented as durable state with approver and timestamp.\n\nAlternatives Considered\n1. Preserve each regional workflow independently. Rejected because enterprise reporting and support would remain fragmented.\n2. Wait for the complete Harbor modernization. Rejected because it couples the Atlas pilot to a broader program timeline.\n3. Duplicate CRM identity into an Atlas master. Rejected because reconciliation and ownership become ambiguous.\n4. Automate all exception outcomes. Rejected because legal and contractual cases require accountable human judgment.\n\nFailure Modes and Anti-patterns\nDo not bypass the stage model with hidden flags, use free-text notes as authoritative workflow state, create local identity records when CRM lookup fails, or treat a dependency owner's silence as approval. Do not delete superseded decisions; mark them as superseded so later teams can understand why the architecture evolved.\n\nOperational Considerations\nSupport teams need correlation identifiers across CRM, Atlas, and Harbor. Retry behavior must be idempotent. Audit history must show stage changes and approvals. Dependency outages require visible degraded-state handling. Runbooks must distinguish data-quality failures from authentication and workflow failures.\n\nValidation\nArchitecture acceptance requires traceability to business requirements, dependency-owner review, QA coverage for failure modes, and client confirmation that the resulting operating process is usable.\n\nRelated Knowledge\nSee Atlas steering records, client readiness updates, implementation specifications, QA/UAT plan, Harbor SSO dependency material, and project meeting transcripts. These sources intentionally overlap so a memory system can reconstruct provenance rather than relying on a single canonical paragraph.'''
    rec['content']=body
    m['document_density']='expanded'; m['approx_sections']=9; m['synthetic_document']=True
    return rec

for source in ['sharepoint','onedrive','confluence']:
    path=ROOT/f'{source}.jsonl'
    rows=[json.loads(x) for x in path.read_text().splitlines() if x.strip()]
    rows=[expand(r,source) for r in rows]
    path.write_text('\n'.join(json.dumps(r,ensure_ascii=False) for r in rows)+'\n')

# Rebuild Atlas combined records and global combined corpus.
source_files=['teams.jsonl','outlook.jsonl','transcript.jsonl','sharepoint.jsonl','onedrive.jsonl','confluence.jsonl']
atlas=[]
for f in source_files:
    atlas += [json.loads(x) for x in (ROOT/f).read_text().splitlines() if x.strip()]
atlas.sort(key=lambda r:r['timestamp'])
(ROOT/'all_records.jsonl').write_text('\n'.join(json.dumps(r,ensure_ascii=False) for r in atlas)+'\n')
global_path=Path('data/all_records.jsonl')
if global_path.exists():
    other=[json.loads(x) for x in global_path.read_text().splitlines() if x.strip()]
    other=[r for r in other if r.get('project_id')!='ATLAS']+atlas
    other.sort(key=lambda r:r['timestamp'])
    global_path.write_text('\n'.join(json.dumps(r,ensure_ascii=False) for r in other)+'\n')

# Helpers for substantial artifacts.
styles=getSampleStyleSheet()
def add_docx(filename,title,sections):
    doc=Document(); h=doc.add_paragraph(); run=h.add_run(title); run.bold=True; run.font.size=Pt(22)
    doc.add_paragraph('Contoso | Atlas Customer Onboarding | Synthetic enterprise project artifact')
    for heading,paras in sections:
        doc.add_heading(heading,level=1)
        for para in paras:
            doc.add_paragraph(para)
    doc.save(ART/filename)

def add_pdf(filename,title,sections):
    story=[Paragraph(title,styles['Title']),Paragraph('Contoso | Atlas Customer Onboarding | Synthetic enterprise project artifact',styles['BodyText']),Spacer(1,14)]
    for i,(heading,paras) in enumerate(sections):
        story.append(Paragraph(heading,styles['Heading2']))
        for para in paras: story += [Paragraph(para,styles['BodyText']),Spacer(1,8)]
        if i and i%3==0: story.append(PageBreak())
    SimpleDocTemplate(str(ART/filename),pagesize=letter,rightMargin=48,leftMargin=48,topMargin=48,bottomMargin=48).build(story)

common=[
('Executive Context',["Atlas addresses the gap between a signed enterprise deal and a stable production handoff. The existing process relies on regional checklists, manual status chasing, and inconsistent definitions of when onboarding is truly complete. The target model establishes a governed enterprise stage taxonomy while allowing approved regional substeps.","Phase 1 deliberately focuses on enterprise onboarding. SMB is deferred because its volume, automation expectations, and exception patterns are materially different. The project is expected to prove the operating model before broadening scope."]),
('Business Requirements',["CRM remains authoritative for customer identity and onboarding status. Atlas orchestrates work but does not create a competing customer master. Users need clear task ownership, blocked-state visibility, auditable exceptions, and client-facing status that does not expose internal-only detail.","High-risk legal, security, and contractual exceptions require human approval. The system must record who approved or rejected an exception, the supporting rationale, and the state transition that followed."]),
('Architecture and Dependencies',["Atlas resolves customer identity through CRM, authenticates pilot users through the Harbor interim SSO bridge, and maintains onboarding orchestration state in Atlas-owned services. Harbor modernization continues independently; Atlas does not wait for the complete Harbor migration.","Nova and Atlas share a dependency on CRM identity quality. The projects do not share workflow state, but duplicate or inconsistent customer records can affect both. Cross-project concerns are escalated through enterprise architecture rather than solved with local copies."]),
('Delivery and Risk',["Key risks include duplicate customer records, regional stage mismatch, Harbor SSO timing, legal exception review, and uneven client readiness. Each risk has an accountable owner and must be reflected in acceptance coverage where it can affect production behavior.","A dependency delay does not automatically move the Atlas milestone. The project manager first evaluates resequencing, degraded-mode options, and scope protection before requesting a steering decision."]),
('Governance and Traceability',["Material scope and architecture decisions are retained with rationale, alternatives, owners, and supersession history. Older assumptions are preserved because they explain later design choices. Current truth should be determined from accepted decisions and later revisions, not simply from the newest timestamp.","The project memory corpus intentionally contains overlapping meeting, email, SharePoint, OneDrive, and Confluence evidence so provenance and conflict resolution can be evaluated."])]

add_docx('business_requirements_document.docx','Atlas Business Requirements Document',common+[
('Detailed Process Requirements',["Onboarding begins only after the commercial handoff contains a valid enterprise customer reference and named onboarding owner. The case moves through Initiated, Discovery, Configuration, Validation, Ready for Production, and Completed. Regions may add substeps but may not redefine the meaning of enterprise stages.","Blocked tasks must carry a reason, owner, next review date, and dependency reference where applicable. Client-visible status is derived from governed state rather than manually authored summaries."]),
('Non-Functional Expectations',["The workflow must be auditable, resilient to retried integration calls, and support least-privilege access. Operational teams need searchable correlation identifiers and enough history to distinguish data-quality, authentication, workflow, and user-action failures."])])

add_docx('solution_design_specification.docx','Atlas Solution Design Specification',common+[
('Component Design',["The solution separates the client portal, onboarding API, workflow engine, notification service, onboarding state store, audit history, and document references. Integration boundaries are explicit so that CRM and Harbor remain owners of their respective capabilities.","The workflow engine treats approvals and external dependency waits as durable states. Retryable failures do not advance the business stage. Notifications are emitted from confirmed transitions and are safe to replay without duplicate client messages."]),
('Data Model',["Core entities include CustomerReference, OnboardingCase, StageTransition, Task, Exception, Approval, Dependency, Participant, and ArtifactReference. Each material transition records actor, timestamp, source, previous state, new state, and correlation identifier."])])

add_docx('uat_and_acceptance_plan.docx','Atlas UAT and Acceptance Plan',common+[
('Test Strategy',["UAT combines process validation with integration and control validation. Scenarios cover standard onboarding, duplicate CRM matches, failed Harbor authentication, regional substeps, high-risk exception approval and rejection, stale sessions, retried integration calls, and client status visibility.","Acceptance evidence is retained against the requirement or decision it validates. A successful happy-path demonstration is not sufficient when a known risk lacks negative-path coverage."]),
('Exit Criteria',["No severity-one defects remain open. All phase-1 business requirements have accepted evidence. Harbor bridge support readiness is confirmed. Client product ownership accepts the enterprise stage model and exception experience. Training and operational runbooks reflect the final rather than prototype workflow."])])

add_docx('operational_readiness_runbook.docx','Atlas Operational Readiness and Support Runbook',common+[
('Support Model',["First-line support classifies incidents into identity/data quality, authentication, workflow state, notification, or client-usage categories. Correlation IDs allow support to follow one onboarding case across Atlas, CRM, and Harbor without exposing unnecessary customer data.","Escalations involving enterprise stage semantics go to product/BA ownership; authentication controls go to Harbor; architecture ambiguity goes to the solution architect and enterprise architecture director."]),
('Recovery Guidance',["Retry integration failures only through supported idempotent operations. Never repair a blocked case by editing database state directly or by creating a substitute CRM identity. Material manual remediation requires an audit note and owner."])])

add_docx('regional_process_and_change_plan.docx','Atlas Regional Process and Change Plan',common+[
('Adoption Approach',["Regional teams receive a mapping from their current terminology to the enterprise stage taxonomy. Training focuses on why common stages matter, where regional substeps are still allowed, and how exception ownership changes.","Change champions collect friction during pilot weeks. Feedback may change UI wording or approved substeps, but cannot silently redefine enterprise stage semantics."]),
('Communications',["Leadership communications emphasize reduced handoff ambiguity and clearer client visibility. Operational communications explain task ownership, escalation, and exception handling. Client communications avoid internal architecture language and focus on predictable status and responsibilities."])])

add_pdf('steering_committee_decision_pack.pdf','Atlas Steering Committee Decision Pack',common+[
('Decisions Requiring Continued Oversight',["Enterprise-only phase 1, CRM authority, Harbor interim SSO, human approval for high-risk exceptions, and the common stage taxonomy are the five anchor decisions. Any proposal that contradicts them requires explicit steering or architecture review."]),
('Leadership Actions',["Maintain scope discipline, resolve cross-project ownership quickly, and ensure client-readiness concerns are reflected in the rollout sequence rather than hidden in implementation notes."])])

add_pdf('client_readiness_assessment.pdf','Atlas Client Readiness Assessment',common+[
('Readiness Dimensions',["Readiness is assessed across process ownership, regional mapping, user access, data quality, training, exception governance, and support escalation. A client can be technically connected while still not being operationally ready."]),
('Pilot Recommendation',["Proceed only when CRM identity quality is within agreed tolerance, Harbor pilot authentication is supportable, regional owners accept the stage mapping, and high-risk exception approvers are named."])])

add_pdf('architecture_decision_compendium.pdf','Atlas Architecture Decision Compendium',common+[
('Rejected Approaches',["A separate Atlas customer master was rejected because it creates reconciliation ambiguity. Waiting for full Harbor completion was rejected because it unnecessarily couples program timelines. Fully independent regional workflows were rejected because they preserve the fragmentation Atlas is intended to remove."]),
('Anti-patterns',["Do not encode business stages as undocumented flags, use free text as authoritative state, bypass approvals to unblock a milestone, or duplicate identities when CRM resolution fails. Preserve superseded decisions with their rationale."])])

add_pdf('pilot_post_implementation_review.pdf','Atlas Pilot Post-Implementation Review',common+[
('Observed Outcomes',["The pilot demonstrates that a common stage model improves shared status discussions, but regional teams need clearer guidance on when a local substep is justified. Exception visibility is valuable to operations because blocked cases now have named owners and explicit reasons."]),
('Follow-up Work',["Improve duplicate-record triage, refine training examples, complete Harbor bridge operational handoff, and carry unresolved regional terminology into the next governance review rather than embedding local exceptions in code."])])

add_pdf('data_and_integration_control_review.pdf','Atlas Data and Integration Control Review',common+[
('Control Objectives',["Prevent duplicate authoritative identities, preserve traceable state transitions, enforce least-privilege access, and make integration failure visible. Controls are designed around ownership boundaries rather than attempting to make Atlas authoritative for every piece of data it consumes."]),
('Control Evidence',["Evidence includes CRM lookup outcomes, Harbor authentication logs, Atlas transition audit records, approval history, retry metrics, and UAT scenarios tied to known integration risks."])])

# Additional CSV artifacts.
with (ART/'decision_register.csv').open('w',newline='') as f:
    w=csv.writer(f); w.writerow(['decision_id','topic','decision','owner','rationale','status'])
    for i,d in enumerate(DECISIONS,1): w.writerow([f'ATLAS-D{i:02d}',*d,'Accepted'])
with (ART/'dependency_register.csv').open('w',newline='') as f:
    w=csv.writer(f); w.writerow(['dependency_id','project','dependency','owner','impact','status'])
    w.writerow(['ATLAS-DEP01','HARBOR','Interim SSO bridge','Rhea Pillai','Pilot authentication','Active'])
    w.writerow(['ATLAS-DEP02','NOVA','Shared CRM identity-quality concern','Adrian Shah','Customer/account consistency','Watching'])

print('Atlas enrichment complete:',len(atlas),'records; artifacts:',len(list(ART.iterdir())))
