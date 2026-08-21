from pathlib import Path
import csv, json, random, textwrap
from datetime import datetime, timedelta, timezone
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

random.seed(42)
ROOT = Path('data')
ROOT.mkdir(exist_ok=True)

PROJECTS = [
 {'id':'ATLAS','name':'Atlas Customer Onboarding','domain':'Customer Operations','goal':'Modernize enterprise customer onboarding from signed deal to production handoff.','members':17,'target':420,'themes':['CRM identity','onboarding stage model','exception handling','Harbor SSO dependency','regional process variance']},
 {'id':'PULSE','name':'Pulse Store Operations Intelligence','domain':'Retail Operations','goal':'Create a unified operational intelligence and alerting layer for regional stores.','members':14,'target':360,'themes':['store telemetry','alert thresholds','regional KPI definitions','data quality','shared data platform capacity']},
 {'id':'NOVA','name':'Nova Sales Opportunity Intelligence','domain':'Sales','goal':'Assist strategic account teams with evidence-backed opportunity recommendations.','members':16,'target':400,'themes':['CRM signals','meeting evidence','recommendation explainability','human review','account permissions']},
 {'id':'HARBOR','name':'Harbor Identity Modernization','domain':'Cybersecurity','goal':'Modernize enterprise identity, provisioning and access governance.','members':19,'target':480,'themes':['SSO migration','RBAC','SCIM gaps','privileged access','Atlas interim SSO']},
 {'id':'ORBIT','name':'Orbit Financial Planning','domain':'Finance','goal':'Replace fragmented planning spreadsheets with rolling forecast and scenario planning.','members':15,'target':340,'themes':['ERP actuals','rolling forecast','scenario planning','shared data platform','business-unit template variance']},
]

ROLES = [
 ('Executive Sponsor','leadership'),('Business Owner','leadership'),('Program Manager','management'),('Project Manager','management'),('Product Owner','product'),
 ('Business Analyst','analysis'),('Business Analyst','analysis'),('Solution Architect','architecture'),('Enterprise Architect','architecture'),('Tech Lead','engineering'),
 ('Backend Engineer','engineering'),('Frontend Engineer','engineering'),('Data Engineer','engineering'),('QA Lead','quality'),('QA Engineer','quality'),
 ('Change Management Lead','change'),('Client Product Owner','client'),('Client SME','client'),('Security SME','sme'),('Data SME','sme')]
FIRST=['Maya','Arjun','Priya','Elena','Noah','Victor','Lena','Samir','Ava','Ishaan','Meera','Jon','Grace','Sofia','Marcus','Nina','Daniel','Rachel','Claire','Omar','Tara','Dev','Leah','Ben','Kavya','Rohan','Julia','Adrian','Neha','Miles','Anika','Ethan','Rhea','Gabriel','Sara','Vikram','Mina','Leo','Farah','Joel']
LAST=['Chen','Rao','Shah','Park','Williams','Lee','Ortiz','Patel','Thompson','Stein','Joshi','Bell','Liu','Mendes','Green','Brooks','Kim','Morgan','Reed','Kapoor','Martin','Davis','Thomas','Nair','Mehta','Wilson','Carter','Singh','Lopez','Nguyen','Evans','Bose','White','King','Iyer','Scott','Fernandez','Young','Pillai','Moore']
SHARED=[
 {'id':'ORG-001','name':'Marcus Green','role':'Enterprise Architecture Director','seniority':'leadership','projects':['ATLAS','PULSE','NOVA','HARBOR','ORBIT'],'sme':['enterprise architecture','architecture governance','cross-project dependencies'],'built':['Contoso architecture review board','shared integration principles'],'responsibilities':['Resolves cross-project architecture conflicts and owns enterprise standards.']},
 {'id':'ORG-002','name':'Samir Patel','role':'Data Platform Architect','seniority':'principal','projects':['PULSE','ORBIT','NOVA'],'sme':['data platform','data quality','shared analytical services'],'built':['Contoso analytical data platform'],'responsibilities':['Owns shared analytical platform capacity, standards, and data-quality guidance.']},
 {'id':'ORG-003','name':'Sofia Mendes','role':'Change Management Lead','seniority':'senior','projects':['ATLAS','PULSE','ORBIT'],'sme':['adoption','training','operating model change'],'built':['Contoso adoption playbook'],'responsibilities':['Advises adoption, training, and operating-model changes across transformation programs.']},
]
DECISIONS={
 'ATLAS':[('phase-1 scope','Phase 1 covers enterprise onboarding only; SMB onboarding is deferred.'),('identity dependency','Atlas pilot uses an interim Harbor SSO bridge and does not wait for full Harbor migration.'),('CRM ownership','CRM remains authoritative for customer identity and onboarding status.'),('exceptions','High-risk onboarding exceptions require human approval.'),('regional model','A common stage taxonomy is mandatory, with region-specific substeps allowed.')],
 'PULSE':[('alert scope','Phase 1 alerts target store operations, not employee-level scoring.'),('latency','15-minute analytical latency is accepted; real-time streaming is deferred.'),('pilot','Midwest region is the first production pilot.'),('data platform','Pulse uses the shared analytical platform governed by Samir Patel.'),('KPI governance','Enterprise KPI definitions override legacy regional scorecard definitions.')],
 'NOVA':[('AI authority','Nova recommends actions but cannot automatically change CRM opportunity stages.'),('evidence','Every recommendation must cite CRM, meeting, or account-plan evidence.'),('pilot cohort','Pilot is limited to strategic account directors.'),('evaluation','Expansion requires human-reviewed quality and hallucination thresholds.'),('permissions','Nova may only surface evidence already visible to the requesting account team.')],
 'HARBOR':[('migration','Applications migrate in waves based on criticality and provisioning readiness.'),('RBAC','Role-based access is baseline, with governed exceptions.'),('Atlas bridge','Harbor provides an interim SSO bridge for the Atlas pilot.'),('legacy apps','Legacy apps without SCIM are not forced into phase 1.'),('privileged access','Privileged accounts require step-up controls before migration.')],
 'ORBIT':[('forecast cadence','Monthly rolling forecast replaces quarterly-only planning.'),('actuals','ERP actuals remain authoritative; manual spreadsheet adjustments are exceptions.'),('platform','Orbit uses the shared analytical platform with Pulse.'),('AI scope','AI may draft variance narratives but cannot change forecasts.'),('scenario planning','Scenario planning is phase 1 for finance leadership, phase 2 for business units.')],
}
RISKS={
 'ATLAS':['duplicate customer records','regional stage mismatch','Harbor SSO dependency','legal exception review','client readiness variance'],
 'PULSE':['missing telemetry','high false positives','regional KPI conflicts','platform capacity contention','store adoption resistance'],
 'NOVA':['inconsistent sales notes','recent-meeting bias','qualification disagreement','CRM permission variance','poor evidence quality'],
 'HARBOR':['legacy provisioning gaps','role mapping ambiguity','Atlas timing pressure','privileged-access control change','directory cleanup'],
 'ORBIT':['template incompatibility','platform capacity contention','ERP close delay','scenario scope expansion','manual adjustment governance'],
}
DEPS={'ATLAS':[('HARBOR','interim SSO bridge'),('NOVA','shared CRM data-quality improvements')],'PULSE':[('ORBIT','shared analytical platform capacity')],'NOVA':[('ATLAS','CRM customer and account identity'),('HARBOR','identity role mappings')],'HARBOR':[('ATLAS','pilot authentication dependency'),('NOVA','account-team authorization model')],'ORBIT':[('PULSE','shared platform compute and batch windows')]}
ARCH={
 'ATLAS':[('CRM','Authoritative customer identity'),('Harbor SSO','Interim authentication bridge'),('Onboarding Orchestrator','Stage and exception workflow'),('Client Portal','Customer-facing status')],
 'PULSE':[('Store Systems','Telemetry and operations events'),('Shared Data Platform','15-minute analytical processing'),('Rules Engine','Enterprise KPI thresholds'),('Ops Console','Regional alerts and investigation')],
 'NOVA':[('CRM','Opportunity and account data'),('Meeting Evidence','Approved transcripts and notes'),('Recommendation Service','Evidence-backed suggestions'),('Seller Workspace','Human review and actions')],
 'HARBOR':[('Enterprise Directory','Identity source'),('Provisioning Layer','SCIM and legacy adapters'),('Policy Engine','RBAC and exceptions'),('SSO Gateway','Application access and Atlas bridge')],
 'ORBIT':[('ERP','Authoritative actuals'),('Shared Data Platform','Finance transformations'),('Forecast Engine','Rolling forecast and scenarios'),('Planning Workspace','Finance review and narratives')],
}
KEY={
 'ATLAS':['Enterprise onboarding only in phase 1','CRM is authoritative for customer identity and onboarding status','Harbor provides an interim SSO bridge'],
 'PULSE':['15-minute analytical latency is accepted','Enterprise KPI definitions override regional legacy definitions','Midwest is first production pilot'],
 'NOVA':['Recommendations cannot automatically change CRM stages','Every recommendation requires evidence','Strategic account directors are initial pilot cohort'],
 'HARBOR':['Applications migrate in waves','RBAC is baseline with governed exceptions','Legacy non-SCIM applications may remain outside phase 1'],
 'ORBIT':['Monthly rolling forecast replaces quarterly-only planning','ERP actuals remain authoritative','AI may draft variance narratives but cannot change forecasts'],
}

used={x['name'] for x in SHARED}; pid=1; people=[]; project_people={}
def new_person(prj, role, cat, idx):
 global pid
 while True:
  name=f'{random.choice(FIRST)} {random.choice(LAST)}'
  if name not in used: used.add(name); break
 theme=prj['themes'][idx%len(prj['themes'])]
 p={'id':f'P{pid:03d}','name':name,'role':role,'seniority':cat,'projects':[prj['id']],'sme':[theme,prj['domain'].lower()],'built':[f"{prj['name']} {theme} deliverable",f"{prj['name']} {role.lower()} playbook"],'responsibilities':[f"Owns {theme} decisions and delivery responsibilities for {prj['name']}."]}
 pid+=1; return p
for prj in PROJECTS:
 ps=[]
 for i,(role,cat) in enumerate(ROLES[:prj['members']]):
  p=new_person(prj,role,cat,i); ps.append(p); people.append(p)
 for s in SHARED:
  if prj['id'] in s['projects']: ps.append(s)
 project_people[prj['id']]=ps
all_people=SHARED+people

start=datetime(2026,1,5,9,0,tzinfo=timezone.utc); end=datetime(2026,6,19,18,0,tzinfo=timezone.utc); days=(end-start).days
src_counter={x:1 for x in ['teams','outlook','transcript','sharepoint','onedrive','confluence']}; records=[]
def stamp(i,total,off):
 d=start+timedelta(days=int((i/max(1,total-1))*days),hours=(i*3+off)%8,minutes=(i*7)%55); return d.strftime('%Y-%m-%dT%H:%M:%SZ')
def add(pid0,src,ts,author,typ,content,thread,parts,meta):
 i=src_counter[src]; src_counter[src]+=1; pref={'teams':'TM','outlook':'OL','transcript':'TR','sharepoint':'SP','onedrive':'OD','confluence':'CF'}[src]
 records.append({'id':f'{pref}-{i:05d}','organization':'Contoso','project_id':pid0,'source':src,'type':typ,'timestamp':ts,'author':author,'thread_id':thread,'participants':parts,'content':content,'metadata':{'fictional':True,**meta}})

for prj in PROJECTS:
 id0=prj['id']; ps=project_people[id0]; names=[x['name'] for x in ps]; target=prj['target']; decs=DECISIONS[id0]; risks=RISKS[id0]
 counts={'teams':int(target*.42),'outlook':int(target*.18),'transcript':int(target*.12),'sharepoint':int(target*.10),'onedrive':int(target*.08)}; counts['confluence']=target-sum(counts.values())
 for src,count in counts.items():
  for i in range(count):
   ts=stamp(i,count,{'teams':0,'outlook':1,'transcript':2,'sharepoint':3,'onedrive':4,'confluence':5}[src]); author=names[(i+{'teams':0,'outlook':2,'transcript':0,'sharepoint':1,'onedrive':3,'confluence':4}[src])%len(names)]; topic,dec=decs[i%5]; risk=risks[i%5]; dep=DEPS[id0][i%len(DEPS[id0])]
   if src=='teams':
    modes=[(f'Decision update on {topic}: {dec} This supersedes earlier working assumptions where they conflict.','decision','high'),(f'Current risk: {risk}. We need impact, owner, and mitigation before the next steering review.','risk','high'),(f'Cross-project dependency on {dep[0]}: {dep[1]}. Do not plan this milestone in isolation.','dependency','high'),('Standup: implementation is progressing, but one downstream dependency remains unconfirmed. No formal date change yet.','status','medium'),('The old document is still being referenced. Use the latest ADR and retain the older version only as historical rationale.','superseded','high'),('Client feedback changes the priority order. Update implementation plan and acceptance criteria, not just weekly status.','business_context','high'),('QA found an edge case that is not release-blocking today but should be captured as a known limitation with owner.','quality','medium'),('We need the SME who designed this area. Their earlier decision exists, but the rationale is not obvious from the latest spec.','people_sme','high'),('Confirm whether this is a local project rule or enterprise standard before propagating it.','governance','medium'),('No decision today. Capturing context because this may explain a later scope change.','context','low'),('Action owner confirmed. Follow-up should reference this thread and steering transcript so the decision trail stays connected.','provenance','medium'),(f'Business owner asked whether {risk} changes the promised outcome. Product and architecture need one answer before client review.','leadership','high')]
    c,mt,sig=modes[i%len(modes)]; add(id0,src,ts,author,'message',c,f'{id0}-teams-{i//7}',names[:8],{'signal':sig,'memory_type':mt})
   elif src=='outlook':
    modes=[(f'Subject: Steering decision on {topic}\n\nWe agreed that {dec} Treat this as authoritative unless a later steering decision supersedes it.','decision'),(f'Subject: Risk escalation\n\nCurrent concern: {risk}. Assess business impact and recovery options before next review.','risk'),('Subject: Client feedback\n\nClient changed a workflow expectation. Product, BA and QA need aligned acceptance criteria before implementation continues.','business_context'),('Subject: Correction to previous update\n\nMy earlier note used an outdated assumption. Latest architecture decision and meeting transcript are authoritative.','correction'),('Subject: SME availability\n\nRequired SME is split across another Contoso program. Resequence work rather than treating silence as approval.','people_sme'),('Subject: Milestone status\n\nCurrent milestone is still achievable. Two dependencies remain open, but no approved date change.','status'),('Subject: Leadership rationale\n\nThis item remains prioritized because of operational/customer impact. Preserve that rationale with the implementation decision.','leadership')]
    c,mt=modes[i%len(modes)]; add(id0,src,ts,author,'email',c,f'{id0}-mail-{i//4}',names[:10],{'signal':'medium' if mt=='status' else 'high','memory_type':mt})
   elif src=='transcript':
    parts=names[:8]; c='\n'.join([f"Meeting: {prj['name']} {'Steering Committee' if i%3==0 else 'Working Session'}",f"Participants: {', '.join(parts)}",'',f'00:03 {parts[0]}: We need to resolve {topic} because it affects the next milestone.',f'00:09 {parts[1]}: Main risk is {risk}; current status report understates it.',f'00:16 {parts[2]}: The older option still works technically but increases operating complexity.',f'00:24 {parts[3]}: Client/business preference is to protect the core outcome even if we resequence scope.',f'00:31 {parts[4]}: Cross-project dependency on {dep[0]} is {dep[1]}.',f'00:38 {parts[5]}: Proposed decision: {dec}',f'00:44 {parts[0]}: Agreed. Mark the prior assumption as superseded, not deleted.',f'00:49 {parts[1]}: BA will update requirements; QA will update acceptance coverage.',f'00:55 {parts[2]}: Keep the rejected approach in the ADR because future teams will ask why we did not use it.'])
    add(id0,src,ts,author,'calendar_recording_transcript',c,f'{id0}-meeting-{i}',parts,{'signal':'high','memory_type':'meeting'})
   elif src=='sharepoint':
    typ=['business_announcement','status_report','steering_decision','client_update'][i%4]
    c={'business_announcement':f"Business Announcement\nProject: {prj['name']}\nGoal: {prj['goal']}\nLeadership priority: protect business outcome while managing {risk}.\nLatest decision on {topic}: {dec}",'status_report':f"Weekly Status Report\nOverall: {'AMBER' if i%5==0 else 'GREEN'}\nPrimary risk: {risk}\nCurrent decision: {dec}\nPoint-in-time snapshot; later steering decisions supersede this report.",'steering_decision':f'Steering Decision\nTopic: {topic}\nDecision: {dec}\nBusiness rationale: reduce operational complexity while preserving required capability.\nStatus: Accepted.','client_update':f'Client Update\nClient feedback reviewed against {topic}.\nAgreed direction: {dec}\nOpen concern: {risk}'}[typ]
    add(id0,src,ts,author,typ,c,f'{id0}-sp-{i}',names[:8],{'signal':'high','memory_type':typ,'version':f'{1+i//4}.{i%4}'})
   elif src=='onedrive':
    typ=['implementation_spec','working_notes','test_plan','process_map'][i%4]
    c={'implementation_spec':f'Implementation Specification\nTopic: {topic}\nImplementation must reflect: {dec}\nKnown constraint: {risk}\nOwner: {author}\nNon-goal: do not generalize local behavior into enterprise policy without review.','working_notes':f'Working Notes\nEarlier option retained for comparison but no longer preferred.\nCurrent assumption: {dec}\nQuestion for SME: impact of {risk}.','test_plan':f'QA Test Plan\nDecision under test: {dec}\nRisk coverage: {risk}\nRequired evidence: positive, negative, regression, and client acceptance scenarios.','process_map':f'Process Map Notes\nCurrent business flow follows decision: {dec}\nException path triggered by: {risk}\nClient and BA signoff required.'}[typ]
    add(id0,src,ts,author,typ,c,f'{id0}-od-{i}',names[:6],{'signal':'high','memory_type':typ,'version':f'{1+i//5}.{i%5}'})
   else:
    typ=['architecture_decision','implementation_decision','anti_pattern','dependency_note'][i%4]
    c={'architecture_decision':f'ADR: {topic}\nDecision: {dec}\nContext: {risk}\nAlternatives: retain prior approach; introduce broader shared dependency.\nRationale: minimize operational complexity while preserving required capability.\nStatus: Accepted.','implementation_decision':f'Implementation Decision\nChosen approach: {dec}\nReason: best fit for current phase and operating constraints.\nKnown risk: {risk}\nSupersedes earlier working assumption.','anti_pattern':f'Anti-pattern\nDo not resolve {risk} by bypassing the accepted decision for {topic}.\nWhy: hidden coupling and harder support/debugging.\nPreferred pattern: follow {dec}','dependency_note':f'Dependency Note\nDependent project: {dep[0]}\nDependency: {dep[1]}\nLocal decision: {dec}\nEscalation path: Enterprise Architecture if timelines conflict.'}[typ]
    add(id0,src,ts,author,typ,c,f'{id0}-cf-{i}',names[:7],{'signal':'high','memory_type':typ,'version':f'{1+i//6}.{i%6}'})

# Write normalized dataset and rich file artifacts
(ROOT/'people.json').write_text(json.dumps(all_people,indent=2))
summary={'organization':'Contoso','timeline':{'start':'2026-01-05','end':'2026-06-19'},'projects':PROJECTS,'record_counts':{p['id']:p['target'] for p in PROJECTS},'total_records':len(records),'total_people':len(all_people),'codebase_included':False}
(ROOT/'organization.json').write_text(json.dumps(summary,indent=2))
(ROOT/'all_records.jsonl').write_text('\n'.join(json.dumps(r,ensure_ascii=False) for r in sorted(records,key=lambda x:x['timestamp'])))

for prj in PROJECTS:
 id0=prj['id']; d=ROOT/id0.lower(); d.mkdir(exist_ok=True); art=d/'artifacts'; art.mkdir(exist_ok=True); ps=project_people[id0]; names=[p['name'] for p in ps]; prows=[r for r in records if r['project_id']==id0]
 (d/'project.json').write_text(json.dumps(prj,indent=2)); (d/'people.json').write_text(json.dumps(ps,indent=2)); (d/'all_records.jsonl').write_text('\n'.join(json.dumps(r,ensure_ascii=False) for r in sorted(prows,key=lambda x:x['timestamp'])))
 for src in ['teams','outlook','transcript','sharepoint','onedrive','confluence']:
  rows=[r for r in prows if r['source']==src]; (d/f'{src}.jsonl').write_text('\n'.join(json.dumps(r,ensure_ascii=False) for r in rows))
 with (art/'project_roster.csv').open('w',newline='',encoding='utf-8') as f:
  w=csv.writer(f); w.writerow(['person_id','name','role','seniority','sme_areas','responsibilities','built'])
  for p in ps: w.writerow([p['id'],p['name'],p['role'],p['seniority'],' | '.join(p.get('sme',[])),' | '.join(p.get('responsibilities',[])),' | '.join(p.get('built',[]))])
 with (art/'risk_register.csv').open('w',newline='',encoding='utf-8') as f:
  w=csv.writer(f); w.writerow(['risk_id','risk','owner','probability','impact','status','mitigation','last_reviewed'])
  for i,risk in enumerate(RISKS[id0],1): w.writerow([f'{id0}-R{i:02d}',risk,names[(i+5)%len(names)],['Medium','High','Medium','Low','High'][i-1],['High','High','Medium','High','Medium'][i-1],['Open','Mitigating','Open','Watching','Mitigating'][i-1],f'Mitigation owned by {names[(i+5)%len(names)]}; review before steering. ',f'2026-0{min(6,i+1)}-{10+i:02d}'])
 boxes=[]; x=40; y=150
 for i,(name,desc) in enumerate(ARCH[id0]):
  boxes += [f'<rect x="{x+i*220}" y="{y}" width="180" height="110" rx="16" fill="white" stroke="#334155" stroke-width="2"/>',f'<text x="{x+90+i*220}" y="{y+35}" text-anchor="middle" font-family="Arial" font-size="16" font-weight="700" fill="#0f172a">{name}</text>']
  for j,line in enumerate(textwrap.wrap(desc,22)[:3]): boxes.append(f'<text x="{x+90+i*220}" y="{y+62+j*17}" text-anchor="middle" font-family="Arial" font-size="12" fill="#475569">{line}</text>')
  if i<3: boxes.append(f'<line x1="{x+180+i*220}" y1="{y+55}" x2="{x+220+i*220}" y2="{y+55}" stroke="#64748b" stroke-width="3" marker-end="url(#arrow)"/>')
 (art/'architecture_overview.svg').write_text(f'''<svg xmlns="http://www.w3.org/2000/svg" width="960" height="420" viewBox="0 0 960 420"><defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#64748b"/></marker></defs><rect width="960" height="420" fill="#f8fafc"/><text x="40" y="55" font-family="Arial" font-size="26" font-weight="700" fill="#0f172a">{prj['name']}</text><text x="40" y="88" font-family="Arial" font-size="14" fill="#475569">Synthetic architecture overview for MCP ingestion and project-memory testing</text>{''.join(boxes)}<text x="40" y="350" font-family="Arial" font-size="13" fill="#64748b">Decision context, ownership, dependencies, and implementation details live in linked project artifacts.</text></svg>''')
 doc=Document(); tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=tp.add_run(f"{prj['name']}\nImplementation Specification"); rr.bold=True; rr.font.size=Pt(20); doc.add_paragraph(f'Project ID: {id0}'); doc.add_paragraph(f"Business goal: {prj['goal']}")
 doc.add_heading('1. Scope and Outcomes',1); doc.add_paragraph(f"This specification captures the current implementation direction for {prj['name']}. It is synthetic but structured like a real enterprise artifact in the {prj['domain']} domain.")
 doc.add_heading('2. Architecture Overview',1); [doc.add_paragraph(f'{n}: {dsc}',style='List Bullet') for n,dsc in ARCH[id0]]
 doc.add_heading('3. Key Decisions',1); [doc.add_paragraph(k,style='List Bullet') for k in KEY[id0]]
 doc.add_heading('4. Dependencies and Constraints',1); [doc.add_paragraph(f'{r}: review through risk register and steering decision trail.',style='List Bullet') for r in RISKS[id0]]
 doc.add_heading('5. Roles and SMEs',1); [doc.add_paragraph(f"{p['name']} - {p['role']}. SME: {', '.join(p.get('sme',[])[:2])}.",style='List Bullet') for p in ps[:10]]
 doc.add_heading('6. Acceptance and Governance',1); doc.add_paragraph('Implementation decisions must be traceable to approved project decisions, requirements, client expectations, or architecture constraints. Superseded decisions are retained as history.'); doc.save(art/'implementation_specification.docx')
 styles=getSampleStyleSheet(); story=[Paragraph(f"{prj['name']} - Steering Committee Pack",styles['Title']),Spacer(1,12),Paragraph(f"<b>Project:</b> {id0} &nbsp;&nbsp; <b>Domain:</b> {prj['domain']}",styles['BodyText']),Paragraph(f"<b>Business goal:</b> {prj['goal']}",styles['BodyText']),Spacer(1,14),Paragraph('Executive Summary',styles['Heading2']),Paragraph('The project remains focused on preserving the intended business outcome while keeping architecture, delivery risk, client expectations, and cross-project dependencies explicit. This pack is synthetic and designed for project-memory ingestion.',styles['BodyText']),Spacer(1,10),Paragraph('Top Risks',styles['Heading2'])]
 table=[['Risk','Owner','Status']]+[[RISKS[id0][i],names[(i+6)%len(names)],['Open','Mitigating','Watching','Open','Mitigating'][i]] for i in range(5)]; t=Table(table,colWidths=[3.2*inch,2*inch,1.1*inch]); t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),.5,colors.grey),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold')])); story += [t,PageBreak(),Paragraph('Architecture and Decision Summary',styles['Heading2'])]
 [story.append(Paragraph(f'<b>{n}</b>: {dsc}',styles['BodyText'])) for n,dsc in ARCH[id0]]; story += [Spacer(1,12),Paragraph('Decision Principles',styles['Heading2'])]; [story.append(Paragraph(f'- {k}',styles['BodyText'])) for k in KEY[id0]]; story += [Spacer(1,12),Paragraph('Leadership Ask',styles['Heading2']),Paragraph('Preserve historical rationale, track cross-project dependencies explicitly, and route material scope or architecture changes through proper governance.',styles['BodyText'])]
 SimpleDocTemplate(str(art/'steering_committee_pack.pdf'),pagesize=letter,rightMargin=45,leftMargin=45,topMargin=45,bottomMargin=45).build(story)

print(json.dumps({'records':len(records),'people':len(all_people),'projects':{p['id']:p['target'] for p in PROJECTS}},indent=2))
